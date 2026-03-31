import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header_footer(doc, title_text):
    # Footer with Page Number
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(f"{title_text} | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    p._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    p._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    p._element.append(fldChar2)

def process_inline_formatting(paragraph, text):
    """Processes **bold** and *italic* in text and adds them to the paragraph."""
    # Split text by bold and italic markers
    # This is a simple regex that handles **bold** and *italic*
    # We use a pattern that captures the markers and the text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('[') and '](' in part:
            # Handle [link](url) - just show label as bold/blue for now
            label = re.search(r'\[(.*?)\]', part).group(1)
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0, 0, 238)
            run.underline = True
        else:
            paragraph.add_run(part)

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    # Section setup
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    title_main = "PCIC RO10 Technical Documentation"
    add_header_footer(doc, title_main)

    in_code_block = False
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip('\n'))
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.5)
            continue

        # Handle Tables
        if "|" in line:
            if not in_table:
                in_table = True
                table_data = []
            
            if re.match(r'^[\s|:-]+$', stripped):
                continue
            
            parts = line.split("|")
            if parts[0].strip() == "": parts = parts[1:]
            if parts and parts[-1].strip() == "": parts = parts[:-1]
            
            cells = [c.strip() for c in parts]
            if cells:
                table_data.append(cells)
            continue
        else:
            if in_table:
                if table_data:
                    columns = max(len(row) for row in table_data)
                    table = doc.add_table(rows=0, cols=columns)
                    table.style = 'Table Grid'
                    for i, row_cells in enumerate(table_data):
                        row = table.add_row()
                        for j, cell_text in enumerate(row_cells):
                            if j < columns:
                                cell = row.cells[j]
                                cell_p = cell.paragraphs[0]
                                process_inline_formatting(cell_p, cell_text)
                                # Shade header row
                                if i == 0:
                                    shading_elm_1 = OxmlElement('w:shd')
                                    shading_elm_1.set(qn('w:fill'), 'D9D9D9')
                                    cell._tc.get_or_add_tcPr().append(shading_elm_1)
                                    cell_p.runs[0].bold = True if cell_p.runs else False
                in_table = False
                table_data = []

        # Handle Headings
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        
        # Handle List Items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, content)
        
        # Handle Info/Alerts (starts with >)
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            # Add a vertical bar logic is hard in docx, just make it italic
            run = p.add_run("Note: ")
            run.bold = True
            process_inline_formatting(p, stripped[2:])
            for run in p.runs:
                run.italic = True

        # Handle Body Text
        elif stripped:
            p = doc.add_paragraph()
            process_inline_formatting(p, stripped)
        else:
            # Empty line -> spacing
            doc.add_paragraph()

    doc.save(docx_path)
    print(f"Successfully generated {docx_path}")

if __name__ == "__main__":
    md_file = "TECHNICAL_DOCUMENTATION.md"
    docx_file = "PCIC_RO10_Technical_Documentation_v1.4.1.docx"
    markdown_to_docx(md_file, docx_file)
